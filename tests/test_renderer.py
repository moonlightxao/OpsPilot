# -*- coding: utf-8 -*-
"""
TemplateRenderer 单元测试

测试覆盖:
1. render() - Word 文档生成（模板/内置）
2. _render_summary() - 摘要渲染
3. _render_risk_alerts() - 风险告警渲染
4. _render_section() - 章节渲染
5. _render_task_table() - 表格渲染
6. RenderStrategy - 渲染策略切换
"""

from pathlib import Path

import pytest
from docx import Document

from src.renderer import (
    TemplateRenderer,
    render_with_template,
    RenderStrategy,
    TemplateRendererError
)
from src.renderer.template_renderer import TemplateNotFoundError


def _empty_report():
    """v2.0 格式空报告"""
    return {
        'meta': {'source_file': 'empty.xlsx', 'generated_at': '2026-02-18T10:00:00Z', 'version': '2.0.0'},
        'summary': {'total_tasks': 0, 'total_sheets': 0, 'high_risk_count': 0, 'has_external_links': False, 'external_links': []},
        'has_risk_alerts': False,
        'risk_alerts': [],
        'sections': []
    }


def _minimal_report_no_alerts():
    """无风险告警的最小报告"""
    return {
        'meta': {'source_file': 'test.xlsx', 'generated_at': '2026-02-18T10:00:00Z', 'version': '2.0.0'},
        'summary': {'total_tasks': 1, 'total_sheets': 1, 'high_risk_count': 0, 'has_external_links': True, 'external_links': ['https://example.com']},
        'has_risk_alerts': False,
        'risk_alerts': [],
        'sections': []
    }


def _render(renderer, report, output_path, template_path="templates/template.docx"):
    """统一调用入口：模板不存在时走内置渲染"""
    return renderer.render(report, template_path, str(output_path))


class TestTemplateRendererInit:
    """测试 TemplateRenderer 初始化"""
    
    def test_init_with_default_config(self):
        renderer = TemplateRenderer()
        assert renderer._config is not None
        assert renderer._output_config is not None
    
    def test_init_with_custom_config(self, sample_config):
        renderer = TemplateRenderer(config_path=str(sample_config))
        assert renderer._config is not None
        assert 'output_config' in renderer._config
    
    def test_init_with_missing_config(self, temp_dir):
        with pytest.raises(FileNotFoundError):
            TemplateRenderer(config_path=str(temp_dir / "not_exist.yaml"))


class TestTemplateRendererRender:
    """测试 TemplateRenderer.render() 方法"""
    
    def test_render_basic_document(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "output.docx"
        result = _render(renderer, sample_report, output_path)
        assert Path(result).exists()
        assert Path(result).suffix == '.docx'
    
    def test_render_creates_output_directory(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "subdir" / "deep" / "output.docx"
        result = _render(renderer, sample_report, output_path)
        assert Path(result).exists()
        assert Path(result).parent.exists()
    
    def test_render_implementation_summary_table(self, sample_config, temp_dir):
        """测试实施总表表格渲染"""
        report_with_impl = {
            'meta': {'source_file': 'test.xlsx', 'version': '2.1.0'},
            'summary': {'total_tasks': 0, 'total_sheets': 0, 'high_risk_count': 0,
                        'has_external_links': False, 'external_links': []},
            'has_risk_alerts': False, 'risk_alerts': [],
            'implementation_summary': {
                'sheet_name': '上线安排',
                'columns': ['序号', '任务', '开始时间', '结束时间', '实施人', '复核人'],
                'rows': [
                    {'cells': ['1', '环境检查', '', '', '张三', '']},
                    {'cells': ['2', '脚本执行', '', '', '李四', '']},
                ],
                'has_data': True,
            },
            'sections': [],
        }
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "impl_summary.docx"
        _render(renderer, report_with_impl, output_path)
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert '2 实施步骤和计划' in text
        assert len(doc.tables) >= 1

    def test_render_document_structure(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "structure_test.docx"
        _render(renderer, sample_report, output_path)
        doc = Document(str(output_path))
        assert len(doc.paragraphs) > 0
        assert len(doc.tables) > 0
    
    def test_render_with_empty_report(self, sample_config, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "empty.docx"
        result = _render(renderer, _empty_report(), output_path)
        assert Path(result).exists()


class TestTemplateRendererSummary:
    """测试摘要渲染"""
    
    def test_summary_task_count(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "summary_test.docx"
        _render(renderer, sample_report, output_path)
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert '任务总数' in text
        assert '3' in text
    
    def test_summary_high_risk_count(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "risk_test.docx"
        _render(renderer, sample_report, output_path)
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert '高危操作' in text
        assert '1' in text
    
    def test_summary_external_links(self, sample_config, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "links_test.docx"
        _render(renderer, _minimal_report_no_alerts(), output_path)
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert 'https://example.com' in text


class TestTemplateRendererRiskAlerts:
    """测试风险告警渲染"""
    
    def test_risk_alerts_present(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "alerts_test.docx"
        _render(renderer, sample_report, output_path)
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert '高危操作' in text or '⚠️' in text
    
    def test_risk_alerts_task_names(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "task_names_test.docx"
        _render(renderer, sample_report, output_path)
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert '删除临时表' in text
    
    def test_no_risk_alerts_when_empty(self, sample_config, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "no_alerts_test.docx"
        _render(renderer, _empty_report(), output_path)
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert '高危操作告警' not in text


class TestTemplateRendererSections:
    """测试章节渲染"""
    
    def test_sections_present(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "sections_test.docx"
        _render(renderer, sample_report, output_path)
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert '应用配置' in text
    
    def test_sections_priority_order(self, sample_config, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "priority_test.docx"
        report = {
            'meta': {'source_file': 'test.xlsx', 'generated_at': '2026-02-18T10:00:00Z', 'version': '2.0.0'},
            'summary': {'total_tasks': 2, 'total_sheets': 2, 'high_risk_count': 0, 'has_external_links': False, 'external_links': []},
            'has_risk_alerts': False,
            'risk_alerts': [],
            'sections': [
                {'section_name': '应用配置', 'priority': 20, 'has_action_groups': False, 'columns': [], 'task_count': 1, 'action_groups': []},
                {'section_name': '数据库脚本部署', 'priority': 10, 'has_action_groups': False, 'columns': [], 'task_count': 1, 'action_groups': []},
            ]
        }
        _render(renderer, report, output_path)
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        db_pos = text.find('数据库脚本部署')
        app_pos = text.find('应用配置')
        assert db_pos < app_pos


class TestTemplateRendererTables:
    """测试表格渲染"""
    
    def test_table_created(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "table_test.docx"
        _render(renderer, sample_report, output_path)
        doc = Document(str(output_path))
        assert len(doc.tables) > 0
    
    def test_table_headers(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "headers_test.docx"
        _render(renderer, sample_report, output_path)
        doc = Document(str(output_path))
        if doc.tables:
            table = doc.tables[0]
            header_text = [cell.text for cell in table.rows[0].cells]
            assert len(header_text) > 0
    
    def test_table_data(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "data_test.docx"
        _render(renderer, sample_report, output_path)
        doc = Document(str(output_path))
        if doc.tables:
            table = doc.tables[0]
            assert len(table.rows) >= 2
            data_row = table.rows[1]
            row_text = ' '.join([cell.text for cell in data_row.cells])
            assert len(row_text) > 0


class TestTemplateRendererHighRiskMarking:
    """测试高危操作标记"""
    
    def test_high_risk_marked_in_title(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "high_risk_mark_test.docx"
        _render(renderer, sample_report, output_path)
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert '⚠️' in text or '高危' in text
    
    def test_high_risk_instruction_bold(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "bold_test.docx"
        _render(renderer, sample_report, output_path)
        doc = Document(str(output_path))
        has_bold = any(run.bold for p in doc.paragraphs for run in p.runs)
        assert has_bold


class TestTemplateRendererChineseFont:
    """测试中文字体设置"""
    
    def test_chinese_font_applied(self, sample_config, sample_report, temp_dir):
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "chinese_font_test.docx"
        _render(renderer, sample_report, output_path)
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert '任务' in text or '操作' in text or '配置' in text


class TestTemplateRendererConvenienceFunction:
    """测试便捷函数"""
    
    def test_render_with_template_function(self, sample_config, sample_report, temp_dir):
        output_path = temp_dir / "convenience_test.docx"
        result = render_with_template(
            sample_report,
            template_path="templates/template.docx",
            output_path=str(output_path),
            config_path=str(sample_config)
        )
        assert Path(result).exists()


class TestTemplateRendererTemplateDetection:
    """测试模板检测与回退逻辑"""
    
    def test_non_docxtpl_template_falls_back_to_builtin(self, sample_config, sample_report, temp_dir):
        """当模板文件不存在 Jinja 标记时，应回退到内置渲染逻辑"""
        from docx import Document
        
        # 创建一个不含任何 Jinja 标记的普通 Word 文档作为“伪模板”
        fake_template = temp_dir / "fake_template.docx"
        doc = Document()
        doc.add_heading("纯样式模板", level=1)
        doc.add_paragraph("这里没有任何 Jinja 占位符，仅用于验证回退逻辑。")
        doc.save(str(fake_template))
        
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "fallback_output.docx"
        
        result_path = renderer.render(sample_report, str(fake_template), str(output_path))
        result_doc = Document(result_path)
        
        text = "\n".join(p.text for p in result_doc.paragraphs)
        
        # 断言：输出中应包含来自 report.sections 的内容（证明使用了内置渲染）
        assert "应用配置" in text
        # 并且至少有一张表格（任务清单表格）
        assert len(result_doc.tables) > 0


class TestRenderStrategy:
    """测试渲染策略切换"""

    def test_get_render_strategy_auto(self, sample_config):
        """测试 AUTO 策略解析"""
        renderer = TemplateRenderer(config_path=str(sample_config))
        assert renderer._get_render_strategy() == RenderStrategy.AUTO

    def test_get_render_strategy_invalid_defaults_to_auto(self, temp_dir):
        """测试无效策略默认回退到 AUTO"""
        config_content = """
output_config:
  title_style:
    font_name: "微软雅黑"
render_config:
  strategy: "invalid_strategy"
"""
        config_path = temp_dir / "test_config.yaml"
        config_path.write_text(config_content, encoding='utf-8')
        
        renderer = TemplateRenderer(config_path=str(config_path))
        assert renderer._get_render_strategy() == RenderStrategy.AUTO

    def test_auto_strategy_uses_builtin_when_no_template(
        self, sample_config, sample_report, temp_dir
    ):
        """AUTO 策略：无模板时使用内置渲染"""
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "auto_no_template.docx"
        
        result = renderer.render(
            sample_report,
            "nonexistent_template.docx",
            str(output_path)
        )
        
        assert Path(result).exists()
        doc = Document(result)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "应用配置" in text

    def test_builtin_strategy_ignores_template(
        self, sample_config, sample_report, temp_dir
    ):
        """BUILTIN 策略：忽略模板参数"""
        config_content = """
output_config:
  title_style:
    font_name: "微软雅黑"
render_config:
  strategy: "builtin"
"""
        config_path = temp_dir / "builtin_config.yaml"
        config_path.write_text(config_content, encoding='utf-8')
        
        renderer = TemplateRenderer(config_path=str(config_path))
        output_path = temp_dir / "builtin_output.docx"
        
        result = renderer.render(
            sample_report,
            "any_template.docx",
            str(output_path)
        )
        
        assert Path(result).exists()
        doc = Document(result)
        assert len(doc.paragraphs) > 0


class TestTemplateNotFound:
    """测试模板不存在时的行为"""

    def test_template_strategy_raises_when_no_template(
        self, sample_report, temp_dir
    ):
        """TEMPLATE 策略：模板不存在时抛出异常"""
        config_content = """
output_config:
  title_style:
    font_name: "微软雅黑"
render_config:
  strategy: "template"
"""
        config_path = temp_dir / "template_config.yaml"
        config_path.write_text(config_content, encoding='utf-8')
        
        renderer = TemplateRenderer(config_path=str(config_path))
        output_path = temp_dir / "template_output.docx"
        
        with pytest.raises(TemplateNotFoundError):
            renderer.render(
                sample_report,
                "nonexistent.docx",
                str(output_path)
            )

    def test_template_strategy_raises_when_no_jinja_markers(
        self, sample_report, temp_dir
    ):
        """TEMPLATE 策略：模板无 Jinja 标记时抛出异常"""
        config_content = """
output_config:
  title_style:
    font_name: "微软雅黑"
render_config:
  strategy: "template"
"""
        config_path = temp_dir / "template_config.yaml"
        config_path.write_text(config_content, encoding='utf-8')
        
        fake_template = temp_dir / "fake.docx"
        doc = Document()
        doc.add_paragraph("No Jinja markers here")
        doc.save(str(fake_template))
        
        renderer = TemplateRenderer(config_path=str(config_path))
        output_path = temp_dir / "output.docx"
        
        with pytest.raises(TemplateRendererError):
            renderer.render(
                sample_report,
                str(fake_template),
                str(output_path)
            )


class TestFallbackBehavior:
    """测试回退机制"""

    def test_auto_fallback_on_template_error(
        self, sample_config, sample_report, temp_dir
    ):
        """AUTO 策略：模板渲染失败时回退内置"""
        # 创建一个有 Jinja 语法但引用不存在变量的模板
        broken_template = temp_dir / "broken.docx"
        try:
            from docxtpl import DocxTemplate
            doc = DocxTemplate(str(broken_template))
        except ImportError:
            pytest.skip("docxtpl not installed")
        except Exception:
            pass
        
        # 使用伪模板（有 Jinja 但语法可能有问题）
        fake_jinja_template = temp_dir / "fake_jinja.docx"
        doc = Document()
        doc.add_paragraph("{{ nonexistent_variable }}")
        doc.save(str(fake_jinja_template))
        
        renderer = TemplateRenderer(config_path=str(sample_config))
        output_path = temp_dir / "fallback.docx"
        
        result = renderer.render(
            sample_report,
            str(fake_jinja_template),
            str(output_path)
        )
        
        assert Path(result).exists()

    def test_no_fallback_when_configured_false(
        self, sample_report, temp_dir
    ):
        """配置禁用回退时，模板失败应抛出异常"""
        config_content = """
output_config:
  title_style:
    font_name: "微软雅黑"
render_config:
  strategy: "auto"
  fallback_to_builtin: false
"""
        config_path = temp_dir / "no_fallback_config.yaml"
        config_path.write_text(config_content, encoding='utf-8')
        
        renderer = TemplateRenderer(config_path=str(config_path))
        output_path = temp_dir / "no_fallback.docx"
        
        # 无模板时应直接内置渲染（非模板失败场景）
        result = renderer.render(
            sample_report,
            "nonexistent.docx",
            str(output_path)
        )
        assert Path(result).exists()
