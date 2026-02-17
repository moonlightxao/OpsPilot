# -*- coding: utf-8 -*-
"""
端到端测试 (E2E Tests)

测试完整的 analyze -> generate 流程
"""

import json
import subprocess
import sys
from pathlib import Path

import pytest


class TestEndToEnd:
    """端到端测试"""
    
    def test_full_workflow_with_cli(self, sample_config, sample_excel, temp_dir):
        """测试完整 CLI 工作流：analyze -> generate"""
        output_json = temp_dir / "report.json"
        output_docx = temp_dir / "实施文档.docx"
        
        # 阶段 1: analyze
        result_analyze = subprocess.run(
            [
                sys.executable, "main.py", "analyze",
                str(sample_excel),
                "--output", str(output_json),
                "--config", str(sample_config)
            ],
            capture_output=True,
            text=True,
            cwd=Path(__file__).parent.parent
        )
        
        # 验证 analyze 成功
        assert result_analyze.returncode == 0, f"Analyze failed: {result_analyze.stderr}"
        assert output_json.exists()
        
        # 验证 report.json 内容
        with open(output_json, 'r', encoding='utf-8') as f:
            report = json.load(f)
        
        assert 'meta' in report
        assert 'summary' in report
        assert 'sections' in report
        
        # 阶段 2: generate
        result_generate = subprocess.run(
            [
                sys.executable, "main.py", "generate",
                str(output_json),
                "--output", str(output_docx),
                "--config", str(sample_config)
            ],
            capture_output=True,
            text=True,
            cwd=Path(__file__).parent.parent
        )
        
        # 验证 generate 成功
        assert result_generate.returncode == 0, f"Generate failed: {result_generate.stderr}"
        assert output_docx.exists()
    
    def test_run_command_full_flow(self, sample_config, sample_excel, temp_dir):
        """测试 run 命令完整流程"""
        output_docx = temp_dir / "完整流程.docx"
        
        result = subprocess.run(
            [
                sys.executable, "main.py", "run",
                str(sample_excel),
                "--output", str(output_docx),
                "--config", str(sample_config),
                "--force"  # 强制执行（跳过高危确认）
            ],
            capture_output=True,
            text=True,
            cwd=Path(__file__).parent.parent
        )
        
        # 验证成功
        assert result.returncode == 0, f"Run failed: {result.stderr}"
        assert output_docx.exists()
        
        # 验证中间文件也生成了
        output_json = output_docx.with_suffix('.json')
        assert output_json.exists()
    
    def test_analyze_with_high_risk_detection(self, sample_config, sample_excel_with_high_risk, temp_dir):
        """测试 analyze 检测高危操作"""
        output_json = temp_dir / "risk_report.json"
        
        result = subprocess.run(
            [
                sys.executable, "main.py", "analyze",
                str(sample_excel_with_high_risk),
                "--output", str(output_json),
                "--config", str(sample_config)
            ],
            capture_output=True,
            text=True,
            cwd=Path(__file__).parent.parent
        )
        
        # 验证输出包含高危告警
        assert "高危操作" in result.stdout or "high_risk" in result.stdout.lower()
        
        # 验证 report.json 标记了高危
        with open(output_json, 'r', encoding='utf-8') as f:
            report = json.load(f)
        
        assert report['summary']['high_risk_count'] > 0
    
    def test_run_blocks_without_force_on_high_risk(self, sample_config, sample_excel_with_high_risk, temp_dir):
        """测试 run 命令在高危操作时阻止执行（无 --force）"""
        output_docx = temp_dir / "blocked.docx"
        
        result = subprocess.run(
            [
                sys.executable, "main.py", "run",
                str(sample_excel_with_high_risk),
                "--output", str(output_docx),
                "--config", str(sample_config)
                # 注意：没有 --force
            ],
            capture_output=True,
            text=True,
            cwd=Path(__file__).parent.parent
        )
        
        # 应该被阻止（非零返回码）
        assert result.returncode != 0
        assert "高危操作" in result.stdout or "force" in result.stdout.lower()
        assert not output_docx.exists()
    
    def test_data_integrity_from_excel_to_word(self, sample_config, sample_excel, temp_dir):
        """测试数据完整性：Excel -> JSON -> Word"""
        from src.parser import ExcelParser
        from src.renderer import WordRenderer
        from docx import Document
        
        # 解析 Excel
        parser = ExcelParser(config_path=str(sample_config))
        report = parser.parse(str(sample_excel))
        
        # 验证任务数量
        assert report['summary']['total_tasks'] == 5  # 3 + 2
        
        # 收集所有任务名
        excel_tasks = set()
        for section in report['sections']:
            for group in section['action_groups']:
                for task in group['tasks']:
                    excel_tasks.add(task['task_name'])
        
        assert '创建用户表' in excel_tasks
        assert '添加索引' in excel_tasks
        assert '删除临时表' in excel_tasks
        assert '部署服务A' in excel_tasks
        assert '升级服务B' in excel_tasks
        
        # 生成 Word
        output_docx = temp_dir / "integrity_test.docx"
        renderer = WordRenderer(config_path=str(sample_config))
        renderer.render(report, str(output_docx))
        
        # 验证 Word 中包含所有任务
        doc = Document(str(output_docx))
        word_text = '\n'.join([p.text for p in doc.paragraphs])
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    word_text += cell.text + '\n'
        
        # 所有任务名应该出现在 Word 中
        for task in excel_tasks:
            assert task in word_text, f"任务 '{task}' 未在 Word 中找到"
    
    def test_priority_order_preserved(self, sample_config, sample_excel, temp_dir):
        """测试优先级顺序在输出中保持"""
        from src.parser import ExcelParser
        from src.renderer import WordRenderer
        from docx import Document
        
        # 解析
        parser = ExcelParser(config_path=str(sample_config))
        report = parser.parse(str(sample_excel))
        
        # 验证 sections 按 priority 排序
        priorities = [s['priority'] for s in report['sections']]
        assert priorities == sorted(priorities)
        
        # 生成 Word
        output_docx = temp_dir / "priority_test.docx"
        renderer = WordRenderer(config_path=str(sample_config))
        renderer.render(report, str(output_docx))
        
        # 验证 Word 中顺序正确
        doc = Document(str(output_docx))
        text = '\n'.join([p.text for p in doc.paragraphs])
        
        # 数据库脚本部署(10) 应该在 上线代码包(15) 之前
        db_pos = text.find('数据库脚本部署')
        code_pos = text.find('上线代码包')
        
        assert db_pos < code_pos, "章节优先级顺序错误"


class TestCLICommands:
    """CLI 命令测试"""
    
    def test_cli_help(self):
        """测试 CLI 帮助信息"""
        result = subprocess.run(
            [sys.executable, "main.py", "--help"],
            capture_output=True,
            text=True,
            cwd=Path(__file__).parent.parent
        )
        
        assert result.returncode == 0
        assert "OpsPilot" in result.stdout
    
    def test_analyze_help(self):
        """测试 analyze 命令帮助"""
        result = subprocess.run(
            [sys.executable, "main.py", "analyze", "--help"],
            capture_output=True,
            text=True,
            cwd=Path(__file__).parent.parent
        )
        
        assert result.returncode == 0
        assert "analyze" in result.stdout.lower() or "分析" in result.stdout
    
    def test_generate_help(self):
        """测试 generate 命令帮助"""
        result = subprocess.run(
            [sys.executable, "main.py", "generate", "--help"],
            capture_output=True,
            text=True,
            cwd=Path(__file__).parent.parent
        )
        
        assert result.returncode == 0
        assert "generate" in result.stdout.lower() or "生成" in result.stdout
    
    def test_run_help(self):
        """测试 run 命令帮助"""
        result = subprocess.run(
            [sys.executable, "main.py", "run", "--help"],
            capture_output=True,
            text=True,
            cwd=Path(__file__).parent.parent
        )
        
        assert result.returncode == 0
        assert "run" in result.stdout.lower() or "完整" in result.stdout
