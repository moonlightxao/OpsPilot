# -*- coding: utf-8 -*-
"""
黄金样本测试：使用 docs/Sample_Files/上线checklist.xlsx 和 实施文档.docx 作为基准
"""
import json
from pathlib import Path

import pytest
from docx import Document

from src.parser import ExcelParser
from src.renderer import TemplateRenderer


class TestGoldenSample:
    """黄金样本端到端测试"""

    @pytest.fixture
    def golden_paths(self):
        base = Path(__file__).parent.parent / "docs" / "Sample_Files"
        return {
            "xlsx": base / "上线checklist.xlsx",
            "docx": base / "实施文档.docx",
            "config": Path(__file__).parent.parent / "config" / "rules.yaml",
        }

    @pytest.fixture
    def temp_output(self, temp_dir):
        return temp_dir / "golden_output.docx"

    @pytest.mark.skipif(
        not (Path(__file__).parent.parent / "docs" / "Sample_Files" / "上线checklist.xlsx").exists(),
        reason="黄金样本 上线checklist.xlsx 不存在（docs/Sample_Files 被 gitignore）",
    )
    def test_golden_sample_e2e(self, golden_paths, temp_output):
        """使用黄金样本 Excel 完整解析并生成 Word，验证数据无遗漏"""
        xlsx = golden_paths["xlsx"]
        config = golden_paths["config"]
        parser = ExcelParser(config_path=str(config))
        report = parser.parse(str(xlsx))

        renderer = TemplateRenderer(config_path=str(config))
        renderer.render(report, "templates/template.docx", str(temp_output))

        assert temp_output.exists(), "Word 输出未生成"

        # 数据完整性：sections 中的任务应全部出现在 Word 中
        excel_tasks = set()
        for s in report.get("sections", []):
            for g in s.get("action_groups", []):
                for t in g.get("tasks", []):
                    cells = t.get("cells", [])
                    if cells:
                        excel_tasks.add(str(cells[0]).strip())

        doc = Document(str(temp_output))
        word_text = "\n".join([p.text for p in doc.paragraphs])
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    word_text += cell.text + "\n"

        missing = [t for t in excel_tasks if t and t not in word_text]
        assert not missing, f"以下任务未在 Word 中找到: {missing}"

    @pytest.mark.skipif(
        not (Path(__file__).parent.parent / "docs" / "Sample_Files" / "上线checklist.xlsx").exists(),
        reason="黄金样本 上线checklist.xlsx 不存在",
    )
    def test_golden_implementation_summary(self, golden_paths):
        """验证黄金样本解析出实施总表"""
        parser = ExcelParser(config_path=str(golden_paths["config"]))
        report = parser.parse(str(golden_paths["xlsx"]))

        impl = report.get("implementation_summary", {})
        assert "implementation_summary" in report
        assert impl.get("has_data") or impl.get("rows") is not None
        # 第一个 Sheet 有数据时应有 rows
        if impl.get("sheet_name"):
            assert "columns" in impl
            assert "rows" in impl
