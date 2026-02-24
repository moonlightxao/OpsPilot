# -*- coding: utf-8 -*-
"""
生成文档与样例文档结构对比测试
依据 PRD 第6节、PM 验收标准逐项比对
"""
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn


def _extract_headings(doc: Document) -> list[str]:
    """提取文档中所有标题文本（按出现顺序）"""
    headings = []
    for p in doc.paragraphs:
        if p.style and p.style.name and "Heading" in p.style.name:
            t = p.text.strip()
            if t:
                headings.append(t)
    return headings


def _extract_table_headers(doc: Document) -> list[list[str]]:
    """提取每个表格的表头行（第一行单元格文本）"""
    result = []
    for t in doc.tables:
        if t.rows:
            headers = [c.text.strip() for c in t.rows[0].cells]
            result.append(headers)
    return result


def _extract_all_text(doc: Document) -> str:
    """提取文档全部文本（段落+表格）"""
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _check_implementation_table_columns(headers: list[str]) -> tuple[bool, list[str]]:
    """检查实施总表是否为标准 6 列，返回 (是否通过, 异常列列表)"""
    expected = {"序号", "任务", "开始时间", "结束时间", "实施人", "复核人"}
    bad = []
    for h in headers:
        h_clean = h.strip()
        if not h_clean:
            continue
        if h_clean in expected:
            continue
        # 异常：Excel 日期序列号、Unnamed 等
        if h_clean.isdigit() or "Unnamed" in h_clean:
            bad.append(h_clean)
        elif h_clean not in expected:
            bad.append(h_clean)
    return len(bad) == 0, bad


# 标准样例应包含的结构（依据 TECH_DESIGN v2 与 PM 验收）
REQUIRED_HEADINGS = [
    "1 原因和目的",
    "1.1 变更应用",
    "1.2 变更原因和目的",
    "1.3 变更影响",
    "2 实施步骤和计划",
    "2.1 详细实施步骤",
    "3 实施后验证计划",
    "4 应急回退措施",
    "5 风险分析和规避措施",
]


class TestTemplateComparison:
    """生成文档与样例文档对比"""

    @pytest.fixture
    def paths(self):
        base = Path(__file__).parent.parent
        sample_docx = base / "docs" / "Sample_Files" / "实施文档.docx"
        output_docx = base / "output" / "实施文档.docx"
        return {"sample": sample_docx, "output": output_docx}

    @pytest.mark.skipif(
        not (Path(__file__).parent.parent / "docs" / "Sample_Files" / "实施文档.docx").exists(),
        reason="样例 实施文档.docx 不存在",
    )
    def test_output_vs_sample_structure(self, paths):
        """
        比对：生成文档 vs 样例文档 的章节结构
        依据 PM 验收标准：第1/2/3～5 部分须存在
        """
        sample = paths["sample"]
        output = paths["output"]
        if not output.exists():
            pytest.skip("output/实施文档.docx 未生成，请先执行 main.py run")

        doc_sample = Document(str(sample))
        doc_output = Document(str(output))

        headings_sample = _extract_headings(doc_sample)
        headings_output = _extract_headings(doc_output)

        defects = []

        for required in REQUIRED_HEADINGS:
            in_sample = any(required in h or h == required for h in headings_sample)
            in_output = any(required in h or h == required for h in headings_output)
            if in_sample and not in_output:
                defects.append(f"【缺失】标准有 '{required}'，生成文档缺失")
            if not in_sample:
                # 标准样例可能结构略不同，以 REQUIRED 为准
                pass

        if defects:
            msg = "生成文档与样例结构不一致：\n" + "\n".join(defects)
            msg += f"\n\n样例标题: {headings_sample}"
            msg += f"\n生成标题: {headings_output}"
            pytest.fail(msg)

    @pytest.mark.skipif(
        not (Path(__file__).parent.parent / "output" / "实施文档.docx").exists(),
        reason="output/实施文档.docx 未生成",
    )
    def test_implementation_table_columns(self, paths):
        """
        比对：实施总表列名须为 序号|任务|开始时间|结束时间|实施人|复核人
        不得出现 46315、Unnamed 等
        """
        output = paths["output"]
        if not output.exists():
            pytest.skip("output/实施文档.docx 未生成")

        doc = Document(str(output))
        table_headers = _extract_table_headers(doc)

        defects = []
        for i, headers in enumerate(table_headers):
            ok, bad = _check_implementation_table_columns(headers)
            if not ok and bad:
                # 可能是实施总表（第一个多列表格）
                if len(headers) >= 6 or any("Unnamed" in h or h.isdigit() for h in headers):
                    defects.append(f"表格{i + 1} 列异常: {bad} | 全部列: {headers}")

        if defects:
            pytest.fail("实施总表列不符合标准（6列：序号|任务|开始时间|结束时间|实施人|复核人）:\n" + "\n".join(defects))

    @pytest.mark.skipif(
        not (Path(__file__).parent.parent / "docs" / "Sample_Files" / "实施文档.docx").exists(),
        reason="样例 实施文档.docx 不存在",
    )
    def test_section_2_structure(self, paths):
        """
        第2部分结构：2 实施步骤和计划 → 2.1 详细实施步骤 → 2.1.1～2.1.N
        不得仅为 2.1 实施总表、2.2 详细步骤
        """
        output = paths["output"]
        if not output.exists():
            pytest.skip("output/实施文档.docx 未生成")

        doc = Document(str(output))
        headings = _extract_headings(doc)

        has_21_detail = any("2.1 详细实施步骤" in h or "2.1 详细" in h for h in headings)
        has_wrong_21 = any("2.1 实施总表" in h for h in headings)
        has_wrong_22 = any("2.2 详细步骤" in h for h in headings)

        defects = []
        if has_wrong_21:
            defects.append("标准应为「2.1 详细实施步骤」，当前为「2.1 实施总表」")
        if has_wrong_22:
            defects.append("标准应为 2.1.1～2.1.N 子节，不得使用「2.2 详细步骤」")
        if not has_21_detail and (has_wrong_21 or has_wrong_22):
            defects.append("第2部分层级与命名不符样例")

        if defects:
            pytest.fail("第2部分结构错误:\n" + "\n".join(defects) + f"\n当前标题: {headings}")
