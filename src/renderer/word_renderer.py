# -*- coding: utf-8 -*-
"""
Word Renderer Module
负责基于 python-docx 的文档渲染

设计原则:
1. 样式驱动：所有样式配置从 rules.yaml 读取
2. 结构化渲染：按 Section -> ActionGroup -> TaskTable 层级
3. 防御性编程：处理空数据和缺失字段
"""

from pathlib import Path
from typing import Any, Optional

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class WordRendererError(Exception):
    """Word 渲染异常基类"""
    pass


class WordRenderer:
    """
    Word 文档渲染器

    职责：
    - 读取 report.json 中间态数据
    - 基于 rules.yaml 的样式配置
    - 生成符合样式的 Word 实施文档
    """

    def __init__(self, config_path: str = "config/rules.yaml"):
        """
        初始化渲染器

        Args:
            config_path: 规则配置文件路径
        """
        self.config_path = Path(config_path)
        self._config: dict = {}
        self._output_config: dict = {}
        self._sheet_column_mapping: dict = {}
        self._default_columns: list = []
        self._action_library: dict = {}

        self._load_config()

    def _load_config(self) -> None:
        """加载并解析规则配置文件"""
        if not self.config_path.exists():
            raise FileNotFoundError(f"配置文件不存在: {self.config_path}")

        with open(self.config_path, 'r', encoding='utf-8') as f:
            self._config = yaml.safe_load(f)

        # 提取输出配置
        self._output_config = self._config.get('output_config', {})
        self._sheet_column_mapping = self._config.get('sheet_column_mapping', {})
        self._default_columns = self._config.get('default_columns', [])
        self._action_library = self._config.get('action_library', {})

    def render(self, report: dict, output_path: str) -> str:
        """
        渲染 Word 文档

        Args:
            report: 符合 report_schema.md 的字典数据
            output_path: 输出文件路径

        Returns:
            实际输出的文件路径
        """
        # 创建新文档
        doc = Document()

        # 设置中文字体（解决中文字体问题）
        self._setup_chinese_font(doc)

        # 渲染文档标题
        meta = report.get('meta', {})
        source_file = meta.get('source_file', '实施文档')
        title_text = f"{Path(source_file).stem} - 实施方案"
        self._add_title(doc, title_text)

        # 渲染摘要信息
        self._render_summary(doc, report.get('summary', {}))

        # 渲染风险告警（如果有）
        risk_alerts = report.get('risk_alerts', [])
        if risk_alerts:
            self._render_risk_alerts(doc, risk_alerts)

        # 渲染各章节
        sections = report.get('sections', [])
        # 按 priority 排序
        sections_sorted = sorted(sections, key=lambda x: x.get('priority', 999))

        for section in sections_sorted:
            self._render_section(doc, section)

        # 确保输出目录存在
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        # 保存文档
        doc.save(str(output_file))

        return str(output_file)

    def _setup_chinese_font(self, doc: Document) -> None:
        """
        设置文档默认中文字体支持

        Args:
            doc: Document 对象
        """
        # 设置文档默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _add_title(self, doc: Document, text: str) -> None:
        """
        添加文档标题

        Args:
            doc: Document 对象
            text: 标题文本
        """
        title_style = self._output_config.get('title_style', {})
        font_name = title_style.get('font_name', '微软雅黑')
        font_size = title_style.get('font_size', 16)

        heading = doc.add_heading(text, level=0)
        run = heading.runs[0]
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = run.font.size  # 保持默认大小或自定义
        run.font.bold = title_style.get('bold', True)

    def _render_summary(self, doc: Document, summary: dict) -> None:
        """
        渲染摘要信息

        Args:
            doc: Document 对象
            summary: 摘要数据
        """
        doc.add_paragraph()  # 空行

        # 添加摘要段落
        p = doc.add_paragraph()
        p.add_run("【摘要信息】").bold = True

        total_tasks = summary.get('total_tasks', 0)
        total_sheets = summary.get('total_sheets', 0)
        high_risk_count = summary.get('high_risk_count', 0)

        doc.add_paragraph(f"• 任务总数：{total_tasks}")
        doc.add_paragraph(f"• 涉及模块：{total_sheets} 个")
        doc.add_paragraph(f"• 高危操作：{high_risk_count} 个")

        # 外部链接
        external_links = summary.get('external_links', [])
        if external_links:
            doc.add_paragraph("• 参考链接：")
            for link in external_links:
                p = doc.add_paragraph(f"  - {link}", style='List Bullet')

        doc.add_paragraph()  # 空行分隔

    def _render_risk_alerts(self, doc: Document, risk_alerts: list) -> None:
        """
        渲染风险告警

        Args:
            doc: Document 对象
            risk_alerts: 风险告警列表
        """
        # 添加醒目的风险提示
        p = doc.add_paragraph()
        run = p.add_run("⚠️ 高危操作告警")
        run.bold = True
        run.font.color.rgb = None  # 可设置为红色

        for alert in risk_alerts:
            sheet_name = alert.get('sheet_name', '未知')
            action_type = alert.get('action_type', '未知')
            task_count = alert.get('task_count', 0)

            p = doc.add_paragraph()
            p.add_run(f"【{sheet_name}】{action_type}").bold = True
            p.add_run(f" - 共 {task_count} 个任务")

            # 列出具体任务
            task_names = alert.get('task_names', [])
            for task_name in task_names:
                doc.add_paragraph(f"  • {task_name}")

        doc.add_paragraph()  # 空行分隔

    def _render_section(self, doc: Document, section: dict) -> None:
        """
        渲染单个章节

        Args:
            doc: Document 对象
            section: 章节数据
        """
        section_name = section.get('section_name', '未知章节')

        # 添加章节标题（二级标题）
        heading = doc.add_heading(section_name, level=1)

        # 渲染各操作组
        action_groups = section.get('action_groups', [])

        for i, action_group in enumerate(action_groups):
            # 操作组之间添加空行（第一个除外）
            if i > 0:
                doc.add_paragraph()

            self._render_action_group(doc, section_name, action_group)

    def _render_action_group(
        self,
        doc: Document,
        section_name: str,
        action_group: dict
    ) -> None:
        """
        渲染操作组（操作说明 + 任务表格）

        Args:
            doc: Document 对象
            section_name: 所属章节名
            action_group: 操作组数据
        """
        action_type = action_group.get('action_type', '未知操作')
        instruction = action_group.get('instruction', f"执行以下{action_type}操作：")
        is_high_risk = action_group.get('is_high_risk', False)
        tasks = action_group.get('tasks', [])

        # 添加操作类型标题（三级标题）
        heading = doc.add_heading(action_type, level=2)
        if is_high_risk:
            # 高危操作添加标记
            run = heading.runs[0]
            run.text = f"⚠️ {action_type}"

        # 添加操作说明
        p = doc.add_paragraph()
        if is_high_risk:
            p.add_run(instruction).bold = True
        else:
            p.add_run(instruction)

        # 渲染任务表格
        if tasks:
            self._render_task_table(doc, section_name, tasks)

    def _render_task_table(
        self,
        doc: Document,
        section_name: str,
        tasks: list
    ) -> None:
        """
        渲染任务清单表格

        Args:
            doc: Document 对象
            section_name: 所属章节名
            tasks: 任务列表
        """
        # 获取该章节应展示的列
        columns = self._get_columns_for_section(section_name)

        # 创建表格
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置表头
        header_cells = table.rows[0].cells
        table_header_style = self._output_config.get('table_header_style', {})

        for i, col_name in enumerate(columns):
            cell = header_cells[i]
            cell.text = col_name

            # 设置表头样式
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = table_header_style.get('bold', True)
                    run.font.name = table_header_style.get('font_name', '微软雅黑')
                    run._element.rPr.rFonts.set(
                        qn('w:eastAsia'),
                        table_header_style.get('font_name', '微软雅黑')
                    )

            # 设置表头背景色
            bg_color = table_header_style.get('background_color', '#D9E2F3')
            self._set_cell_shading(cell, bg_color)

        # 填充数据行
        body_style = self._output_config.get('body_style', {})

        for task in tasks:
            row_cells = table.add_row().cells
            raw_data = task.get('raw_data', {})

            for i, col_name in enumerate(columns):
                cell = row_cells[i]
                # 从 raw_data 中获取对应的值
                value = raw_data.get(col_name, '')

                # 如果 raw_data 中没有，尝试从核心字段获取
                if not value:
                    value = self._get_value_from_task(task, col_name)

                cell.text = str(value) if value else ''

                # 设置正文样式
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = body_style.get('font_name', '宋体')
                        run._element.rPr.rFonts.set(
                            qn('w:eastAsia'),
                            body_style.get('font_name', '宋体')
                        )

    def _get_columns_for_section(self, section_name: str) -> list:
        """
        获取指定章节应展示的列名

        Args:
            section_name: 章节/Sheet 名称

        Returns:
            列名列表
        """
        if section_name in self._sheet_column_mapping:
            return self._sheet_column_mapping[section_name].get('columns', [])
        return self._default_columns

    def _get_value_from_task(self, task: dict, col_name: str) -> str:
        """
        从任务数据中获取指定列的值

        支持从核心字段和 raw_data 中查找

        Args:
            task: 任务数据
            col_name: 列名

        Returns:
            单元格值
        """
        # 列名到核心字段的映射
        field_mapping = {
            '任务名': 'task_name',
            '任务名称': 'task_name',
            '执行人': 'executor',
            '部署单元': 'deploy_unit',
            '外部链接': 'external_link',
        }

        core_field = field_mapping.get(col_name)
        if core_field and core_field in task:
            return task.get(core_field, '')

        return ''

    def _set_cell_shading(self, cell, color: str) -> None:
        """
        设置单元格背景色

        Args:
            cell: 表格单元格
            color: 颜色值（如 #D9E2F3）
        """
        # 移除 # 前缀
        if color.startswith('#'):
            color = color[1:]

        # 创建 shading 元素
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)

        # 应用到单元格
        cell._tc.get_or_add_tcPr().append(shading_elm)


# 便捷函数
def render_report(
    report: dict,
    output_path: str,
    config_path: str = "config/rules.yaml"
) -> str:
    """
    渲染 Word 文档的便捷函数

    Args:
        report: 符合 report_schema.md 的字典数据
        output_path: 输出文件路径
        config_path: 规则配置文件路径

    Returns:
        实际输出的文件路径
    """
    renderer = WordRenderer(config_path)
    return renderer.render(report, output_path)
