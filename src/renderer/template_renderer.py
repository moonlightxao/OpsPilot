# -*- coding: utf-8 -*-
"""
Template Renderer Module
基于 docxtpl 的模板渲染器

设计原则:
1. 模板驱动：所有样式定义在 Word 模板中，确保与样板一致
2. Jinja2 语法：使用 {% for %} 循环和 {{ }} 变量
3. 防御性编程：处理空数据和缺失字段
4. 双轨渲染：支持模板渲染与内置渲染两种模式
"""

from enum import Enum
from pathlib import Path
from typing import Any, Optional

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class RenderStrategy(Enum):
    """渲染策略枚举"""
    AUTO = "auto"
    TEMPLATE = "template"
    BUILTIN = "builtin"


class TemplateRendererError(Exception):
    """模板渲染异常基类"""
    pass


class TemplateNotFoundError(TemplateRendererError):
    """模板文件不存在异常"""
    pass


class TemplateRenderer:
    """
    模板渲染器
    
    职责：
    - 基于 docxtpl 模板引擎渲染 Word 文档
    - 支持 Jinja2 循环和条件语法
    - 支持三种渲染策略：AUTO、TEMPLATE、BUILTIN
    """
    
    def __init__(self, config_path: str = "config/rules.yaml"):
        """
        初始化渲染器
        
        Args:
            config_path: 规则配置文件路径
        """
        self.config_path = Path(config_path)
        self._config: dict = {}
        self._output_config: dict = {}
        self._render_config: dict = {}
        
        self._load_config()
    
    def _load_config(self) -> None:
        """加载并解析规则配置文件"""
        if not self.config_path.exists():
            raise FileNotFoundError(f"配置文件不存在: {self.config_path}")
        
        with open(self.config_path, 'r', encoding='utf-8') as f:
            self._config = yaml.safe_load(f)
        
        self._output_config = self._config.get('output_config', {})
        self._render_config = self._config.get('render_config', {})
    
    def render(self, report: dict, template_path: str, output_path: str) -> str:
        """
        渲染 Word 文档
        
        Args:
            report: 符合 report_schema.md v2.0 的字典数据
            template_path: 模板文件路径
            output_path: 输出文件路径
            
        Returns:
            实际输出的文件路径
        """
        strategy = self._get_render_strategy()
        
        if strategy == RenderStrategy.BUILTIN:
            return self._render_builtin(report, output_path)
        
        if strategy == RenderStrategy.TEMPLATE:
            return self._render_with_template_strict(report, template_path, output_path)
        
        return self._render_auto(report, template_path, output_path)
    
    def _get_render_strategy(self) -> RenderStrategy:
        """获取渲染策略"""
        strategy_str = self._render_config.get('strategy', 'auto')
        try:
            return RenderStrategy(strategy_str)
        except ValueError:
            return RenderStrategy.AUTO
    
    def _render_auto(self, report: dict, template_path: str, output_path: str) -> str:
        """AUTO 策略：优先模板，失败回退内置"""
        template_file = Path(template_path)
        
        if template_file.exists() and self._is_docxtpl_template(template_file):
            try:
                return self._render_with_template(report, template_file, output_path)
            except Exception as e:
                if self._render_config.get('fallback_to_builtin', True):
                    import warnings
                    warnings.warn(f"模板渲染失败，回退到内置渲染: {e}")
                    return self._render_builtin(report, output_path)
                raise
        
        return self._render_builtin(report, output_path)
    
    def _render_with_template_strict(
        self,
        report: dict,
        template_path: str,
        output_path: str
    ) -> str:
        """TEMPLATE 策略：严格使用模板，不回退"""
        template_file = Path(template_path)
        
        if not template_file.exists():
            raise TemplateNotFoundError(f"模板文件不存在: {template_path}")
        
        if not self._is_docxtpl_template(template_file):
            raise TemplateRendererError(
                f"模板文件不包含 Jinja2 标记，无法用于模板渲染: {template_path}"
            )
        
        return self._render_with_template(report, template_file, output_path)
    
    def _is_docxtpl_template(self, template_path: Path) -> bool:
        """检查是否是有效的 docxtpl 模板
        
        仅当文档中实际包含 Jinja2/docxtpl 标记时才视为有效模板，
        否则回退到内置渲染逻辑，避免误将纯 Word 文档当作模板。
        """
        try:
            from docxtpl import DocxTemplate
            doc = DocxTemplate(str(template_path))
            
            def _has_jinja(text: Any) -> bool:
                if not text:
                    return False
                s = str(text)
                return "{{" in s or "{%" in s
            
            # 使用 get_docx() 获取 Document 对象（docx 属性在新版本返回 None）
            docx_doc = doc.get_docx()
            
            # 检查段落中的 Jinja2 标记
            for paragraph in docx_doc.paragraphs:
                if _has_jinja(paragraph.text):
                    return True
            
            # 检查表格单元格中的 Jinja2 标记
            for table in docx_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if _has_jinja(cell.text):
                            return True
            
            # 未检测到任何 Jinja2 标记，视为非模板文档
            return False
        except ImportError:
            # docxtpl 未安装，回退到内置渲染
            return False
        except Exception:
            # 模板解析错误（如 Jinja 语法错误），回退到内置渲染
            return False
    
    def _render_with_template(
        self,
        report: dict,
        template_path: Path,
        output_path: str
    ) -> str:
        """使用 docxtpl 模板渲染，失败时回退到内置渲染"""
        try:
            from docxtpl import DocxTemplate
        except ImportError:
            raise TemplateRendererError(
                "docxtpl 未安装，请运行: pip install docxtpl"
            )

        try:
            doc = DocxTemplate(str(template_path))
            context = self._prepare_context(report)
            doc.render(context)

            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_file))

            return str(output_file)
        except Exception as e:
            # 模板渲染失败（如 Jinja2 语法错误），回退到内置渲染
            import warnings
            warnings.warn(f"模板渲染失败，回退到内置渲染: {e}")
            return self._render_builtin(report, output_path)
    
    def _render_builtin(self, report: dict, output_path: str) -> str:
        """
        内置渲染逻辑（兼容 v2.0 数据结构）
        
        使用 python-docx 直接生成文档，无需模板文件
        输出结构符合 PRD 5. 文档输出规范，包含 5 个部分
        """
        doc = Document()
        self._setup_chinese_font(doc)
        
        # 渲染标题
        meta = report.get('meta', {})
        source_file = meta.get('source_file', '实施文档')
        title_text = f"{Path(source_file).stem} - 实施方案"
        self._add_title(doc, title_text)
        
        # ===== 第1部分：原因和目的 =====
        doc.add_heading('1 原因和目的', level=1)
        
        # 1.1 变更应用
        doc.add_heading('1.1 变更应用', level=2)
        app_name = meta.get('application_name', '（待填写）')
        doc.add_paragraph(app_name)
        
        # 1.2 变更原因和目的
        doc.add_heading('1.2 变更原因和目的', level=2)
        change_reason = meta.get('change_reason', '（待填写）')
        doc.add_paragraph(change_reason)
        
        # 1.3 变更影响
        doc.add_heading('1.3 变更影响', level=2)
        change_impact = meta.get('change_impact', '（待填写）')
        doc.add_paragraph(change_impact)
        
        # 渲染摘要
        self._render_summary(doc, report.get('summary', {}))
        
        # 渲染风险告警
        if report.get('has_risk_alerts', False):
            self._render_risk_alerts(doc, report.get('risk_alerts', []))
        
        # ===== 第2部分：实施步骤和计划 =====
        doc.add_heading('2 实施步骤和计划', level=1)
        impl_summary = report.get('implementation_summary', {})
        if impl_summary.get('has_data', False):
            self._render_implementation_summary_table(
                doc, impl_summary.get('columns', []), impl_summary.get('rows', [])
            )
            doc.add_paragraph()
        
        # 2.1 详细实施步骤
        doc.add_heading('2.1 详细实施步骤', level=3)
        sections = report.get('sections', [])
        sections_sorted = sorted(sections, key=lambda x: x.get('priority', 999))

        for idx, section in enumerate(sections_sorted, start=1):
            self._render_section(doc, section, idx)
        
        # ===== 第3部分：实施后验证计划 =====
        doc.add_heading('3 实施后验证计划', level=1)
        doc.add_paragraph('（待填写）')
        
        # ===== 第4部分：应急回退措施 =====
        doc.add_heading('4 应急回退措施', level=1)
        doc.add_paragraph('（待填写）')
        
        # ===== 第5部分：风险分析和规避措施 =====
        doc.add_heading('5 风险分析和规避措施', level=1)
        doc.add_paragraph('（待填写）')
        
        # 保存文档
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file))
        
        return str(output_file)
    
    def _prepare_context(self, report: dict) -> dict:
        """准备模板渲染上下文"""
        impl_summary = report.get('implementation_summary', {})
        if not impl_summary:
            impl_summary = {
                'sheet_name': '',
                'columns': [],
                'rows': [],
                'has_data': False
            }
        # 为 docxtpl 模板填充：实施总表固定 6 列
        if impl_summary.get('has_data'):
            cols = impl_summary.get('columns', [])
            impl_summary = dict(impl_summary)
            impl_summary['columns'] = list(cols) + [''] * max(0, 6 - len(cols))
            padded_rows = []
            for row in impl_summary.get('rows', []):
                cells = list(row.get('cells', []))
                padded_rows.append({'cells': cells + [''] * max(0, 6 - len(cells))})
            impl_summary['rows'] = padded_rows
        meta = dict(report.get('meta', {}))
        meta.setdefault('application_name', '')
        meta.setdefault('change_reason', '')
        meta.setdefault('change_impact', '')
        source_file = meta.get('source_file', '实施文档')
        title = f"{Path(source_file).stem} - 实施方案"
        return {
            'title': title,
            'meta': meta,
            'summary': report.get('summary', {}),
            'has_risk_alerts': report.get('has_risk_alerts', False),
            'risk_alerts': report.get('risk_alerts', []),
            'implementation_summary': impl_summary,
            'sections': report.get('sections', [])
        }
    
    def _setup_chinese_font(self, doc: Document) -> None:
        """设置文档默认中文字体支持"""
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _add_title(self, doc: Document, text: str) -> None:
        """添加文档标题"""
        title_style = self._output_config.get('title_style', {})
        font_name = title_style.get('font_name', '微软雅黑')
        
        heading = doc.add_heading(text, level=0)
        run = heading.runs[0]
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.bold = title_style.get('bold', True)
    
    def _render_summary(self, doc: Document, summary: dict) -> None:
        """渲染摘要信息"""
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run("【摘要信息】").bold = True
        
        doc.add_paragraph(f"• 任务总数：{summary.get('total_tasks', 0)}")
        doc.add_paragraph(f"• 涉及模块：{summary.get('total_sheets', 0)} 个")
        doc.add_paragraph(f"• 高危操作：{summary.get('high_risk_count', 0)} 个")
        
        # 外部链接
        if summary.get('has_external_links', False):
            doc.add_paragraph("• 参考链接：")
            for link in summary.get('external_links', []):
                doc.add_paragraph(f"  - {link}", style='List Bullet')
        
        doc.add_paragraph()
    
    def _render_risk_alerts(self, doc: Document, risk_alerts: list) -> None:
        """渲染风险告警"""
        p = doc.add_paragraph()
        run = p.add_run("⚠️ 高危操作告警")
        run.bold = True
        
        for alert in risk_alerts:
            sheet_name = alert.get('sheet_name', '未知')
            action_type = alert.get('action_type', '未知')
            task_count = alert.get('task_count', 0)
            
            p = doc.add_paragraph()
            p.add_run(f"【{sheet_name}】{action_type}").bold = True
            p.add_run(f" - 共 {task_count} 个任务")
            
            for task_name in alert.get('task_names', []):
                doc.add_paragraph(f"  • {task_name}")
        
        doc.add_paragraph()
    
    def _render_section(self, doc: Document, section: dict, index: int) -> None:
        """渲染单个章节

        Args:
            doc: Document 对象
            section: 章节数据
            index: 章节序号（从1开始）
        """
        section_name = section.get('section_name', '未知章节')

        # 生成带序号的标题：2.1.1 应用配置
        heading_text = f"2.1.{index} {section_name}"
        heading = doc.add_heading(heading_text, level=4)

        # 移除斜体（Heading 4 默认可能是斜体）
        for run in heading.runs:
            run.italic = False

        action_groups = section.get('action_groups', [])
        columns = section.get('columns', [])
        
        for i, action_group in enumerate(action_groups):
            if i > 0:
                doc.add_paragraph()
            self._render_action_group(doc, columns, action_group)
    
    def _render_action_group(
        self,
        doc: Document,
        columns: list,
        action_group: dict
    ) -> None:
        """渲染操作组"""
        action_type = action_group.get('action_type', '未知操作')
        instruction = action_group.get('instruction', f"执行以下{action_type}操作：")
        is_high_risk = action_group.get('is_high_risk', False)
        tasks = action_group.get('tasks', [])

        # 操作类型说明（使用普通段落，加粗显示）
        p = doc.add_paragraph()
        run = p.add_run(f"【{action_type}】")
        run.bold = True
        if is_high_risk:
            run = p.add_run(" ⚠️ 高危操作")
            run.bold = True
            run.font.color.rgb = None  # 保持默认颜色
        
        # 操作说明
        p = doc.add_paragraph()
        if is_high_risk:
            p.add_run(instruction).bold = True
        else:
            p.add_run(instruction)
        
        # 任务表格
        if tasks and columns:
            self._render_task_table(doc, columns, tasks)
    
    def _render_implementation_summary_table(
        self, 
        doc: Document, 
        columns: list, 
        rows: list
    ) -> None:
        """渲染实施总表表格（第2章主体表格）"""
        if not columns or not rows:
            return
        self._render_task_table(doc, columns, rows)
    
    def _render_task_table(
        self, 
        doc: Document, 
        columns: list, 
        tasks: list
    ) -> None:
        """渲染任务清单表格"""
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头样式
        table_header_style = self._output_config.get('table_header_style', {})
        header_cells = table.rows[0].cells
        
        for i, col_name in enumerate(columns):
            cell = header_cells[i]
            cell.text = col_name
            
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = table_header_style.get('bold', True)
                    run.font.name = table_header_style.get('font_name', '微软雅黑')
                    run._element.rPr.rFonts.set(
                        qn('w:eastAsia'),
                        table_header_style.get('font_name', '微软雅黑')
                    )
            
            bg_color = table_header_style.get('background_color', '#D9E2F3')
            self._set_cell_shading(cell, bg_color)
        
        # 数据行
        body_style = self._output_config.get('body_style', {})
        
        for task in tasks:
            cells = task.get('cells', [])
            row_cells = table.add_row().cells
            
            for i in range(min(len(cells), len(columns))):
                cell = row_cells[i]
                cell.text = str(cells[i]) if cells[i] else ''
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = body_style.get('font_name', '宋体')
                        run._element.rPr.rFonts.set(
                            qn('w:eastAsia'),
                            body_style.get('font_name', '宋体')
                        )
    
    def _set_cell_shading(self, cell, color: str) -> None:
        """设置单元格背景色"""
        if color.startswith('#'):
            color = color[1:]
        
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)


def render_with_template(
    report: dict,
    template_path: str = "templates/template.docx",
    output_path: str = "output/实施文档.docx",
    config_path: str = "config/rules.yaml"
) -> str:
    """
    使用模板渲染 Word 文档的便捷函数
    
    Args:
        report: 符合 report_schema.md v2.0 的字典数据
        template_path: 模板文件路径
        output_path: 输出文件路径
        config_path: 规则配置文件路径
        
    Returns:
        实际输出的文件路径
    """
    renderer = TemplateRenderer(config_path)
    return renderer.render(report, template_path, output_path)
