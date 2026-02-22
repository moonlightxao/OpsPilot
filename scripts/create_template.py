# -*- coding: utf-8 -*-
"""
创建 template.docx 模板文件
包含 docxtpl 兼容的 Jinja2 占位符

docxtpl 特殊标签：
- {%p ... %} : 段落级循环/条件
- {%tr ... %} : 表格行级循环
- {%tc ... %} : 表格单元格级循环
- {{ var }} : 普通变量
- {{r var }} : RichText 变量
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


def set_cell_shading(cell, color: str) -> None:
    """设置单元格背景色"""
    if color.startswith('#'):
        color = color[1:]
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_run_font(run, font_name: str, east_asia: bool = True) -> None:
    """设置运行字体"""
    run.font.name = font_name
    if east_asia:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def create_template():
    """创建模板文档"""
    doc = Document()

    # 设置默认中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ========== 文档标题 ==========
    title = doc.add_heading('{{ title }}', level=0)
    if title.runs:
        set_run_font(title.runs[0], '微软雅黑')

    doc.add_paragraph()  # 空行

    # ========== 摘要信息 ==========
    summary_heading = doc.add_paragraph()
    run = summary_heading.add_run('【摘要信息】')
    run.bold = True

    p1 = doc.add_paragraph()
    p1.add_run('• 任务总数：{{ summary.total_tasks }}')

    p2 = doc.add_paragraph()
    p2.add_run('• 涉及模块：{{ summary.total_sheets }} 个')

    p3 = doc.add_paragraph()
    p3.add_run('• 高危操作：{{ summary.high_risk_count }} 个')

    # 外部链接（条件 + 循环）
    p_links = doc.add_paragraph()
    p_links.add_run('{% if summary.has_external_links %}• 参考链接：{% endif %}')

    # 链接循环 - 使用 {%p for %} 创建新段落
    # 但这样会比较复杂，简化处理：假设链接数量有限，直接显示
    doc.add_paragraph('{% for link in summary.external_links %}  - {{ link }}\n{% endfor %}')

    doc.add_paragraph()  # 空行

    # ========== 风险告警区块（条件渲染） ==========
    # docxtpl 使用 {%p if %} 进行段落级条件
    p_risk_title = doc.add_paragraph()
    p_risk_title.add_run('{%p if has_risk_alerts %}')
    p_risk_title.add_run('⚠️ 高危操作告警')

    # 风险告警循环
    # 使用段落级循环
    p_alert = doc.add_paragraph()
    p_alert.add_run('{%p for alert in risk_alerts %}')

    p_alert_content = doc.add_paragraph()
    run_alert = p_alert_content.add_run('【{{ alert.sheet_name }}】{{ alert.action_type }}')
    run_alert.bold = True
    p_alert_content.add_run(' - 共 {{ alert.task_count }} 个任务')

    # 任务名循环
    p_task_names = doc.add_paragraph()
    p_task_names.add_run('{% for task_name in alert.task_names %}  • {{ task_name }}\n{% endfor %}')

    p_alert_end = doc.add_paragraph()
    p_alert_end.add_run('{%p endfor %}')

    p_risk_end = doc.add_paragraph()
    p_risk_end.add_run('{%p endif %}')

    doc.add_paragraph()  # 空行

    # ========== 2.1 实施总表（条件渲染） ==========
    # 实施总表表格：预定义最大 8 列，表头与数据行使用 Jinja2 循环
    impl_start = doc.add_paragraph()
    impl_start.add_run('{%p if implementation_summary.has_data %}')
    impl_heading = doc.add_heading('2.1 实施总表', level=1)
    # 表格：8 列（与常见上线安排列数匹配），动态填充
    impl_table = doc.add_table(rows=2, cols=8)
    impl_table.style = 'Table Grid'
    impl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = impl_table.rows[0].cells
    for i in range(8):
        cell = header_cells[i]
        cell.text = '{%tc if implementation_summary.columns[' + str(i) + '] %}{{ implementation_summary.columns[' + str(i) + '] }}{%tc endif %}'
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                set_run_font(run, '微软雅黑')
        set_cell_shading(cell, 'D9E2F3')
    data_cells = impl_table.rows[1].cells
    data_cells[0].text = '{%tr for row in implementation_summary.rows %}{{ row.cells[0] }}'
    for i in range(1, 7):
        data_cells[i].text = '{{ row.cells[' + str(i) + '] }}'
    data_cells[7].text = '{{ row.cells[7] }}{%tr endfor %}'
    for cell in data_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, '宋体')
    impl_end = doc.add_paragraph()
    impl_end.add_run('{%p endif %}')

    doc.add_paragraph()  # 空行

    # ========== 2.2 详细步骤 - 章节循环 ==========
    # 使用 {%p for section in sections %} 进行段落级循环
    p_section_start = doc.add_paragraph()
    p_section_start.add_run('{%p for section in sections %}')

    # 章节标题（二级标题）
    section_title = doc.add_heading('{{ section.section_name }}', level=1)

    # 操作组循环
    p_action_start = doc.add_paragraph()
    p_action_start.add_run('{%p for action_group in section.action_groups %}')

    # 操作类型标题（三级标题）
    # 高危操作特殊标记
    action_title = doc.add_heading('{{ action_group.action_type }}', level=2)

    # 操作说明
    p_instruction = doc.add_paragraph()
    p_instruction.add_run('{{ action_group.instruction }}')

    # ========== 任务表格 ==========
    # 使用 {%tr for %} 进行表格行级循环
    # 创建模板表格：1行表头 + 1行数据模板（会被循环复制）

    # 表格需要预先定义好列数，这里用最大可能的列数
    # 动态列需要使用 {%tc for %} 但比较复杂
    # 简化方案：固定5列，第一行作为表头

    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头行
    header_cells = table.rows[0].cells
    for i in range(5):
        cell = header_cells[i]
        # 使用 {%tc if %} 条件渲染表头
        cell.text = '{%tc if section.columns[' + str(i) + '] %}{{ section.columns[' + str(i) + '] }}{%tc endif %}'
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                set_run_font(run, '微软雅黑')
        set_cell_shading(cell, 'D9E2F3')

    # 数据行模板 - 使用 {%tr for %} 循环
    # 在表格的第一列放入循环开始标签
    data_cells = table.rows[1].cells

    # 第一列包含循环标签和第一个数据
    data_cells[0].text = '{%tr for task in action_group.tasks %}{{ task.cells[0] }}'
    # 中间列
    for i in range(1, 4):
        data_cells[i].text = '{{ task.cells[' + str(i) + '] }}'
    # 最后一列包含数据 + 循环结束标签
    data_cells[4].text = '{{ task.cells[4] }}{%tr endfor %}'

    # 设置数据行样式
    for cell in data_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, '宋体')

    # 操作组循环结束
    p_action_end = doc.add_paragraph()
    p_action_end.add_run('{%p endfor %}')

    # 章节循环结束
    p_section_end = doc.add_paragraph()
    p_section_end.add_run('{%p endfor %}')

    # 保存模板
    output_path = Path('d:/code/OpsPilot/templates/template.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f'模板已创建: {output_path}')
    return str(output_path)


if __name__ == '__main__':
    create_template()
